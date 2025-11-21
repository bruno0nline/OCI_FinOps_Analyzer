
import os
import csv
from datetime import datetime

from docx import Document
from docx.shared import Pt

DAYS = int(os.getenv("METRICS_DAYS", "30"))
HOME = os.path.expanduser("~")

CSV_FILE = os.path.join(
    HOME, f"Relatorio_CPU_Memoria_media_{DAYS}d_multi_region.csv"
)
DOCX_FILE = os.path.join(
    HOME, f"Relatorio_FinOps_CPU_Mem_{DAYS}d_multi_region.docx"
)

# ===== Config de custo (estimativa) =====
# Valores padrão aproximados, baseados na lista pública de preços da OCI.
# Ajuste via variáveis de ambiente para cada cliente/região:
#   OCI_COST_OCPU_HOUR  (ex: 0.70 para R$/h)
#   OCI_COST_MEM_GB_HOUR (ex: 0.03 para R$/h por GB)
#   OCI_COST_CURRENCY    (ex: BRL, USD, etc.)
CURRENCY = os.getenv("OCI_COST_CURRENCY", "BRL")
COST_OCPU_HOUR = float(os.getenv("OCI_COST_OCPU_HOUR", "0.70"))
COST_MEM_GB_HOUR = float(os.getenv("OCI_COST_MEM_GB_HOUR", "0.03"))
DAYS_IN_MONTH = 30


def to_float(v):
    try:
        return float(v)
    except Exception:
        return None


def estimate_monthly_cost(ocpus, mem_gb):
    if ocpus is None and mem_gb is None:
        return None
    oc = ocpus or 0
    mem = mem_gb or 0
    hourly = oc * COST_OCPU_HOUR + mem * COST_MEM_GB_HOUR
    return hourly * 24 * DAYS_IN_MONTH


def format_money(value):
    if value is None:
        return "N/A"
    return f"{CURRENCY} {value:,.2f}"


def build_downsize_text(row):
    name = row["instance_name"]
    region = row["region"]
    compartment = row["compartment"]
    shape = row["shape"]

    ocpus = to_float(row["ocpus"])
    mem_gb = to_float(row["memory_gb"])
    cpu_mean = to_float(row["cpu_mean_percent"])
    cpu_p95 = to_float(row["cpu_p95_percent"])
    mem_mean = to_float(row["mem_mean_percent"])
    mem_p95 = to_float(row["mem_p95_percent"])
    baseline_percent = row.get("baseline_percent", "Desativada")
    burstable = row.get("burstable_enabled", "NO")

    linhas = []
    linhas.append(f"Instância: {name}")
    linhas.append(f"Região/Compartimento: {region} / {compartment}")
    linhas.append(f"Shape atual: {shape} | OCPUs: {ocpus} | Memória: {mem_gb} GB")
    linhas.append(
        f"Uso médio de CPU: {cpu_mean}% (p95: {cpu_p95}%) | "
        f"Uso médio de Memória: {mem_mean}% (p95: {mem_p95}%)"
    )
    linhas.append(
        f"Instância expansível (burstable): {burstable} | Linha de base: {baseline_percent}"
    )

    target_ocpus = ocpus
    target_mem = mem_gb

    if ocpus and cpu_mean is not None:
        if cpu_mean < 5:
            target_ocpus = max(1, int(round(ocpus * 0.25)))
        elif cpu_mean < 10:
            target_ocpus = max(1, int(round(ocpus * 0.5)))
        elif cpu_mean < 20:
            target_ocpus = max(1, ocpus - 1)

        if target_ocpus < ocpus:
            linhas.append(
                f"➡ Recomenda-se avaliar a redução de CPU de {ocpus} para ~{target_ocpus} OCPUs,"
                f" mantendo monitoramento após o ajuste."
            )

    if mem_gb and mem_mean is not None and mem_mean < 40:
        target_mem = max(1, int(round(mem_gb * 0.7)))
        linhas.append(
            f"➡ Uso de memória estável abaixo de 40%. Avaliar reduzir memória de"
            f" {mem_gb} GB para aproximadamente {target_mem} GB."
        )

    # Cálculo de economia estimada
    custo_atual = estimate_monthly_cost(ocpus, mem_gb)
    custo_sugerido = estimate_monthly_cost(target_ocpus, target_mem)
    if custo_atual is not None and custo_sugerido is not None and custo_sugerido < custo_atual:
        economia = custo_atual - custo_sugerido
        linhas.append(
            f"💰 Estimativa de economia mensal: {format_money(economia)} "
            f"(de {format_money(custo_atual)} para {format_money(custo_sugerido)})."
        )

    if len(linhas) == 4:
        linhas.append("➡ Oportunidade de redução identificada, mas requer análise manual detalhada.")

    return "\n".join(linhas)


def build_upscale_text(row):
    name = row["instance_name"]
    region = row["region"]
    compartment = row["compartment"]
    shape = row["shape"]

    ocpus = to_float(row["ocpus"])
    mem_gb = to_float(row["memory_gb"])
    cpu_mean = to_float(row["cpu_mean_percent"])
    cpu_p95 = to_float(row["cpu_p95_percent"])
    mem_mean = to_float(row["mem_mean_percent"])
    mem_p95 = to_float(row["mem_p95_percent"])
    baseline_percent = row.get("baseline_percent", "Desativada")
    burstable = row.get("burstable_enabled", "NO")

    linhas = []
    linhas.append(f"Instância: {name}")
    linhas.append(f"Região/Compartimento: {region} / {compartment}")
    linhas.append(f"Shape atual: {shape} | OCPUs: {ocpus} | Memória: {mem_gb} GB")
    linhas.append(
        f"Uso médio de CPU: {cpu_mean}% (p95: {cpu_p95}%) | "
        f"Uso médio de Memória: {mem_mean}% (p95: {mem_p95}%)"
    )
    linhas.append(
        f"Instância expansível (burstable): {burstable} | Linha de base: {baseline_percent}"
    )

    target_ocpus = ocpus
    target_mem = mem_gb

    if ocpus and (cpu_p95 or 0) > 80:
        target_ocpus = ocpus + max(1, int(round(ocpus * 0.5)))
        linhas.append(
            f"➡ CPU próxima de saturação (p95 > 80%). Avaliar aumento de OCPUs de"
            f" {ocpus} para ~{target_ocpus} OCPUs ou mudança para forma maior."
        )

    if mem_gb and (mem_p95 or 0) > 85:
        target_mem = int(round(mem_gb * 1.3))
        linhas.append(
            f"➡ Memória próxima de saturação (p95 > 85%). Avaliar aumento de memória de"
            f" {mem_gb} GB para ~{target_mem} GB."
        )

    if burstable == "YES" and baseline_percent == "12.5%":
        linhas.append(
            "➡ Instância expansível com linha de base 12,5% e alto uso."
            " Para cargas críticas, considerar converter para instância regular"
            " (sem burst) com OCPUs dedicadas."
        )

    # Cálculo de aumento de custo estimado
    custo_atual = estimate_monthly_cost(ocpus, mem_gb)
    custo_sugerido = estimate_monthly_cost(target_ocpus, target_mem)
    if custo_atual is not None and custo_sugerido is not None and custo_sugerido > custo_atual:
        aumento = custo_sugerido - custo_atual
        linhas.append(
            f"💰 Estimativa de aumento de custo mensal: {format_money(aumento)} "
            f"(de {format_money(custo_atual)} para {format_money(custo_sugerido)})."
        )

    if len(linhas) == 4:
        linhas.append("➡ Oportunidade de aumento identificada, mas requer análise manual detalhada.")

    return "\n".join(linhas)


def build_burstable_only_text(row):
    name = row["instance_name"]
    region = row["region"]
    compartment = row["compartment"]
    shape = row["shape"]

    ocpus = to_float(row["ocpus"])
    cpu_mean = to_float(row["cpu_mean_percent"])
    cpu_p95 = to_float(row["cpu_p95_percent"])
    burstable = row.get("burstable_enabled", "NO")
    baseline_percent = row.get("baseline_percent", "Desativada")

    if burstable == "YES":
        return None

    if cpu_mean is None:
        return None

    if cpu_mean < 15 and (cpu_p95 or 0) < 60:
        linhas = []
        linhas.append(f"Instância: {name}")
        linhas.append(f"Região/Compartimento: {region} / {compartment}")
        linhas.append(f"Shape atual: {shape} | OCPUs: {ocpus}")
        linhas.append(f"Uso médio de CPU: {cpu_mean}% (p95: {cpu_p95}%)")
        if cpu_mean < 8:
            alvo = "12,5%"
        else:
            alvo = "50%"
        linhas.append(
            "➡ A instância apresenta baixa utilização de CPU com picos moderados."
            f" Se a forma suportar, considerar habilitar instância expansível com"
            f" linha de base de {alvo}, mantendo o mesmo número de OCPUs provisionadas."
        )
        return "\n".join(linhas)

    return None


def main():
    if not os.path.exists(CSV_FILE):
        raise SystemExit(f"Arquivo CSV não encontrado: {CSV_FILE}")

    with open(CSV_FILE, newline="") as f:
        reader = csv.DictReader(f)
        rows = list(reader)

    downsizes = []
    upscales = []
    burstables = []

    for row in rows:
        rec = row.get("finops_recommendation", "KEEP")

        if rec.startswith("DOWNSIZE"):
            downsizes.append(build_downsize_text(row))
        elif rec == "UPSCALE":
            upscales.append(build_upscale_text(row))
        else:
            txt = build_burstable_only_text(row)
            if txt:
                burstables.append(txt)

    doc = Document()

    title = doc.add_heading(
        f"Relatório FinOps – CPU e Memória (janela de {DAYS} dias)", level=0
    )
    title.alignment = 1

    p_info = doc.add_paragraph()
    p_info.add_run(
        f"Gerado em {datetime.now().strftime('%d/%m/%Y %H:%M')} a partir das métricas "
        "históricas de CPU e memória da OCI."
    ).italic = True

    p_cost = doc.add_paragraph()
    p_cost.add_run(
        "Valores de custo estimados com base em preços públicos de OCPU e memória. "
        "Ajuste os parâmetros OCI_COST_OCPU_HOUR, OCI_COST_MEM_GB_HOUR e "
        "OCI_COST_CURRENCY conforme a região/contrato do cliente."
    ).font.size = Pt(8)

    doc.add_paragraph()

    if downsizes:
        doc.add_heading("1. Recomendações de Redução (Downsize)", level=1)
        for bloco in downsizes:
            para = doc.add_paragraph()
            para.style = "Normal"
            run = para.add_run(bloco)
            run.font.size = Pt(10)
            doc.add_paragraph()

    if upscales:
        doc.add_heading("2. Recomendações de Aumento (Upscale)", level=1)
        for bloco in upscales:
            para = doc.add_paragraph()
            run = para.add_run(bloco)
            run.font.size = Pt(10)
            doc.add_paragraph()

    if burstables:
        doc.add_heading("3. Oportunidades para Instâncias Expansíveis (Burstable)", level=1)
        for bloco in burstables:
            para = doc.add_paragraph()
            run = para.add_run(bloco)
            run.font.size = Pt(10)
            doc.add_paragraph()

    if not (downsizes or upscales or burstables):
        doc.add_paragraph(
            "Nenhuma recomendação automática foi gerada. As métricas indicam ambiente estável "
            "ou requerem análise manual mais detalhada."
        )

    doc.save(DOCX_FILE)

    print("✅ Relatório DOCX gerado com sucesso:")
    print(f"   {DOCX_FILE}")


if __name__ == "__main__":
    main()
