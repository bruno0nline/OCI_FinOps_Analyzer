import os
import csv
from docx import Document

DEFAULT_DAYS = 30
DAYS = int(os.getenv("METRICS_DAYS", DEFAULT_DAYS))

HOME = os.path.expanduser("~")
CSV_PATH = os.path.join(HOME, f"Relatorio_CPU_Memoria_media_{DAYS}d_multi_region.csv")
DOCX_PATH = os.path.join(HOME, f"Relatorio_FinOps_TOP5_{DAYS}d.docx")

def load_rows():
    with open(CSV_PATH, newline="") as f:
        return list(csv.DictReader(f))

def get_top5(rows):
    savings, costs = [], []

    for r in rows:
        try:
            value = float(r.get("monthly_savings_brl", 0))
        except:
            continue

        rec = r.get("finops_recommendation", "")
        if value <= 0:
            continue

        if rec.startswith("DOWNSIZE") or "BURSTABLE" in rec:
            savings.append((r, value))
        elif rec == "UPSCALE":
            costs.append((r, value))

    return (
        sorted(savings, key=lambda x: x[1], reverse=True)[:5],
        sorted(costs, key=lambda x: x[1], reverse=True)[:5],
    )

def generate():
    rows = load_rows()
    top_save, top_cost = get_top5(rows)

    doc = Document()
    doc.add_heading("Relatório Executivo – Top 5 FinOps (OCI)", 0)

    doc.add_heading("Top 5 – Maior Economia Potencial", 1)
    for r, v in top_save:
        doc.add_paragraph(f"{r['instance_name']} – {r['finops_recommendation']} – R$ {v:,.2f}")

    doc.add_heading("Top 5 – Maior Impacto de Aumento", 1)
    for r, v in top_cost:
        doc.add_paragraph(f"{r['instance_name']} – UPSCALE – R$ {v:,.2f}")

    doc.save(DOCX_PATH)
    print(f"Relatório Top 5 gerado: {DOCX_PATH}")

if __name__ == "__main__":
    generate()
